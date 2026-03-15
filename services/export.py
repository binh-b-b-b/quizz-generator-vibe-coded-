from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor


def format_results(record: dict) -> dict:
    """
    Shared formatting — returns clean structured data
    used by both PDF and DOCX exporters.
    """
    return {
        "title": f"Quiz Results — {record['topic']}",
        "meta": f"Difficulty: {record['difficulty']} | Type: {record['type']} | Date: {record['date'][:10]}",
        "score_line": f"Score: {record['score']} / {record['total']} ({record['percentage']}%)",
        "answers": record.get("answers", []),
    }


def export_pdf(record: dict) -> bytes:
    """
    Generate a PDF summary of a completed quiz.
    Returns raw bytes to send as a file download.
    """
    data = format_results(record)
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()

    correct_style = ParagraphStyle("correct", parent=styles["Normal"], textColor=colors.HexColor("#16a34a"), fontSize=11)
    wrong_style   = ParagraphStyle("wrong",   parent=styles["Normal"], textColor=colors.HexColor("#dc2626"), fontSize=11)
    meta_style    = ParagraphStyle("meta",    parent=styles["Normal"], textColor=colors.HexColor("#6b7280"), fontSize=10)
    exp_style     = ParagraphStyle("exp",     parent=styles["Normal"], textColor=colors.HexColor("#374151"), fontSize=10, leftIndent=12)

    story = [
        Paragraph(data["title"], styles["Title"]),
        Paragraph(data["meta"], meta_style),
        Spacer(1, 8),
        Paragraph(data["score_line"], styles["Heading2"]),
        Spacer(1, 16),
        Paragraph("Answer Review", styles["Heading2"]),
        Spacer(1, 8),
    ]

    for i, a in enumerate(data["answers"], 1):
        icon = "✓" if a["is_correct"] else "✗"
        style = correct_style if a["is_correct"] else wrong_style
        story.append(Paragraph(f"{i}. {icon}  {a['question']}", style))
        story.append(Paragraph(f"Your answer: {a['user_answer']}", exp_style))
        if not a["is_correct"]:
            story.append(Paragraph(f"Correct: {a['correct_answer']}", exp_style))
        story.append(Paragraph(f"Explanation: {a['explanation']}", exp_style))
        story.append(Spacer(1, 10))

    doc.build(story)
    return buffer.getvalue()


def export_docx(record: dict) -> bytes:
    """
    Generate a .docx summary of a completed quiz.
    Returns raw bytes to send as a file download.
    """
    data = format_results(record)
    doc = Document()

    doc.add_heading(data["title"], level=1)
    doc.add_paragraph(data["meta"]).runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    doc.add_paragraph(data["score_line"]).runs[0].bold = True
    doc.add_heading("Answer Review", level=2)

    for i, a in enumerate(data["answers"], 1):
        icon = "✓" if a["is_correct"] else "✗"
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {icon}  {a['question']}")
        run.bold = True
        run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a) if a["is_correct"] else RGBColor(0xdc, 0x26, 0x26)

        doc.add_paragraph(f"Your answer: {a['user_answer']}", style="List Bullet")
        if not a["is_correct"]:
            doc.add_paragraph(f"Correct: {a['correct_answer']}", style="List Bullet")
        doc.add_paragraph(f"Explanation: {a['explanation']}", style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()